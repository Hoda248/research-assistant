import streamlit as st

# --- 1. PAGE CONFIGURATION & SESSION INITIALIZATION ---
st.set_page_config(page_title="My Research Assistant", page_icon="🧠", layout="wide")

session_defaults = {
    "user_email": None,
    "logged_in": False,
    "role": "user",
    "profile_loaded": False,
    "name": "",
    "keywords":[],
    "authors": "",
    "feed_results": [],
    "discovery_keywords":[],
    "discovery_ran": False,
    "current_page": "Dashboard"
}
for key, value in session_defaults.items():
    if key not in st.session_state:
        st.session_state[key] = value

from Bio import Entrez
from supabase import create_client, Client
from datetime import datetime
import google.generativeai as genai
from docx import Document
from io import BytesIO
import urllib.parse

# --- 2. GLOBAL CSS: PASTEL TURQUOISE & HORIZONTAL ALIGNMENT ---
st.markdown("""
<style>
    /* Global Page Style */
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif !important;
        background-color: #FDFBF7 !important; /* Soft Cream */
        color: #004D40 !important; /* Dark Teal */
    }

    /* THE FIX: Force Streamlit elements to align horizontally inside this container */
    .tag-container {
        display: flex !important;
        flex-wrap: wrap !important;
        gap: 8px !important;
        margin-bottom: 20px !important;
    }

    /* Target the internal Streamlit div to prevent full-width stacking */
    .tag-container > div {
        width: auto !important;
        display: inline-block !important;
    }

    /* Keyword Tags Styling (Pastel Turquoise) */
    .tag-btn div[data-testid="stButton"] > button {
        background-color: #E0F7FA !important; /* Pastel Turquoise */
        color: #004D40 !important;
        border: 1px solid #004D40 !important;
        border-radius: 20px !important;
        padding: 4px 15px !important;
        
        /* Layout Fixes */
        white-space: nowrap !important; /* No text wrapping */
        width: max-content !important;
        min-width: max-content !important;
        display: inline-block !important;
    }

    .tag-btn div[data-testid="stButton"] > button:hover {
        background-color: #EF5350 !important; /* Red for delete */
        color: white !important;
        border-color: #EF5350 !important;
    }

    /* Save/Remove Button Logic */
    .save-btn-wrapper div[data-testid="stButton"] > button {
        background-color: #FFFFFF !important;
        color: #004D40 !important;
        border: 1px solid #004D40 !important;
        white-space: nowrap !important;
    }

    .remove-btn-wrapper div[data-testid="stButton"] > button {
        background-color: #B2EBF2 !important; /* Darker Turquoise */
        color: #004D40 !important;
        border: 1px solid #004D40 !important;
        white-space: nowrap !important;
    }
</style>
""", unsafe_allow_html=True)

# --- 4. GLOBAL HEADER ---
st.markdown("""
    <div style="text-align: center; margin-top: 1rem; margin-bottom: 2.5rem;">
        <h1 style="font-family: 'Lora', serif; color: #004D40; font-size: 3rem; font-weight: 600; margin-bottom: 0;">My Research Assistant</h1>
        <hr style="border: 0; height: 1.5px; background-color: #004D40; margin-top: 15px; margin-bottom: 15px; max-width: 60%; margin-left: auto; margin-right: auto; opacity: 0.6;">
    </div>
""", unsafe_allow_html=True)

# --- SUPABASE DATABASE CONNECTION ---
@st.cache_resource
def init_db_client():
    url = st.secrets.get("SUPABASE_URL")
    key = st.secrets.get("SUPABASE_KEY")
    if not url or not key:
        return None
    try:
        return create_client(url, key)
    except Exception:
        return None

supabase = init_db_client()
if not supabase:
    st.error("🚨 **Database Configuration Error:** Missing or invalid `SUPABASE_URL` / `SUPABASE_KEY` in Streamlit secrets.")
    st.stop()

# --- AUTHENTICATION & SESSION MANAGEMENT ---
def load_user_profile(email):
    try:
        res = supabase.table('users').select('*').eq('email', email).execute()
        return res.data[0] if res.data else None
    except Exception as e:
        st.error(f"Failed to load profile: {e}")
        return None

def log_audit_trail(email):
    try:
        supabase.table('login_history').insert({'user_email': email}).execute()
    except Exception:
        pass

# --- AI LOGIC ---
def get_safe_model():
    api_key = st.secrets.get("AI_API_KEY")
    if not api_key: 
        return None
    try:
        genai.configure(api_key=api_key)
        available =[m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        targets =['models/gemini-1.5-flash', 'gemini-1.5-flash', 'models/gemini-pro']
        for t in targets:
            if t in available: return genai.GenerativeModel(t)
        return genai.GenerativeModel(available[0]) if available else None
    except Exception:
        return None

def generate_ai_summary(abstract):
    model = get_safe_model()
    if not model: return "API Key Configuration Error. Please verify the AI_API_KEY in Streamlit secrets."
    prompt = f"""Summarize this abstract into 4 structured points for a brain researcher:
    1. OBJECTIVE: Main research question.
    2. METHODOLOGY: Tools used (EEG, animal models, fMRI, etc).
    3. FINDINGS: Primary results/mechanisms.
    4. SIGNIFICANCE: Implications for neuroscience/psychopathology.
    Abstract: {abstract}"""
    try:
        return model.generate_content(prompt).text
    except Exception as e: return f"AI Processing Error: {str(e)}"

# --- PUBMED TOOLS ---
def get_summaries(kw_list, author_search, days, logic="OR"):
    if not kw_list and not author_search: return[]
    query = ""
    if kw_list: 
        query += "(" + f" {logic} ".join([f'"{k}"[Title]' for k in kw_list]) + ")"
    if author_search:
        if query: query += " AND "
        query += f"{author_search}[Author]"
    try:
        Entrez.email = st.secrets.get("ADMIN_EMAIL", "admin@example.com")
        h = Entrez.esearch(db="pubmed", term=query, retmax=20, reldate=days)
        ids = Entrez.read(h)["IdList"]
        h.close()
        if not ids: return[]
        h = Entrez.esummary(db="pubmed", id=",".join(ids))
        res = Entrez.read(h)
        h.close()
        return res
    except Exception: 
        return[]

def fetch_abstract(pmid):
    try:
        Entrez.email = st.secrets.get("ADMIN_EMAIL", "admin@example.com")
        h = Entrez.efetch(db="pubmed", id=pmid, retmode="xml")
        recs = Entrez.read(h)
        h.close()
        article = recs['PubmedArticle'][0]['MedlineCitation']['Article']
        return " ".join(article['Abstract']['AbstractText']) if 'Abstract' in article else "Abstract not available."
    except Exception: 
        return "Error retrieving abstract data."

# --- DB HELPERS ---
def toggle_reading_list(pmid, title, journal, authors, date):
    email = st.session_state.user_email
    res = supabase.table('reading_list').select('id').eq('pmid', pmid).eq('user_email', email).execute()
    if res.data:
        supabase.table('reading_list').delete().eq('pmid', pmid).eq('user_email', email).execute()
        return "removed"
    else:
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        data = {
            "user_email": email, "pmid": pmid, "title": title, "journal": journal, 
            "authors": authors, "date": date, "notes": "", "last_edited": now
        }
        supabase.table('reading_list').insert(data).execute()
        return "saved"

def export_notes_to_word():
    doc = Document()
    doc.add_heading(f"My Notebook - {datetime.now().strftime('%Y-%m-%d')}", 0)
    email = st.session_state.user_email
    
    doc.add_heading("Literature Notes", level=1)
    res_notes = supabase.table('reading_list').select('title, notes, authors, date').neq('notes', '').eq('user_email', email).execute()
    for item in res_notes.data:
        year = item['date'][:4] if item['date'] else "n.d."
        doc.add_heading(f"{item['authors']} ({year}). {item['title']}.", level=2)
        doc.add_paragraph(item['notes'])
    
    doc.add_heading("General Research Notes & Thoughts", level=1)
    res_gen = supabase.table('general_notes').select('content, date').eq('user_email', email).order('date', desc=True).execute()
    for item in res_gen.data:
        p = doc.add_paragraph()
        p.add_run(f"[{item['date']}] ").bold = True
        p.add_run(item['content'])
        
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- AUTHENTICATION PORTAL ---
if not st.session_state.logged_in:
    st.markdown("<p style='text-align: center; color: #4A6B58; font-size: 1.2rem; margin-bottom: 40px; margin-top: -10px;'>Neuroscience & Psychopathology Literature Management</p>", unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        tab_login, tab_reg = st.tabs(["Login", "New Investigator Registration"])
        
        with tab_login:
            with st.form("login_form"):
                log_email = st.text_input("Email:")
                if st.form_submit_button("Access Workspace", type="primary", use_container_width=True):
                    cleaned_email = log_email.lower().strip()
                    if cleaned_email:
                        res = supabase.table('users').select('*').eq('email', cleaned_email).execute()
                        if res.data:
                            st.session_state.user_email = cleaned_email
                            st.session_state.logged_in = True
                            log_audit_trail(cleaned_email)
                            st.rerun()
                        else:
                            st.error("Email not found. Please register first.")
                    else:
                        st.error("Please enter a valid email.")
                                
        with tab_reg:
            with st.form("reg_form"):
                reg_name = st.text_input("Investigator Name:")
                reg_email = st.text_input("Email:")
                
                if st.form_submit_button("Register Profile", type="primary", use_container_width=True):
                    cleaned_email = reg_email.lower().strip()
                    if "@" in cleaned_email and reg_name.strip():
                        check_res = supabase.table('users').select('email').eq('email', cleaned_email).execute()
                        if check_res.data:
                            st.error("This email is already registered.")
                        else:
                            new_user = {
                                "email": cleaned_email, 
                                "name": reg_name.strip(),
                                "password_hash": "", 
                                "keywords": "", 
                                "authors": "",
                                "api_key": ""
                            }
                            supabase.table('users').insert(new_user).execute()
                            st.session_state.user_email = cleaned_email
                            st.session_state.logged_in = True
                            log_audit_trail(cleaned_email)
                            st.rerun()
                    else:
                        st.error("Complete all required fields with a valid email.")
    st.stop()

# --- LOAD USER STATE ---
if st.session_state.logged_in and not st.session_state.profile_loaded:
    prof = load_user_profile(st.session_state.user_email)
    if prof:
        st.session_state.name = prof.get('name', '')
        st.session_state.keywords =[k for k in prof.get('keywords', '').split(",") if k] if prof.get('keywords') else[]
        st.session_state.authors = prof.get('authors', '')
    st.session_state.profile_loaded = True

email = st.session_state.user_email

# --- TOP NAVIGATION BAR ---
nav_options =["Dashboard", "Active Tracking", "Literature Discovery", "Reading Room", "My Notebook", "User Guide", "Settings"]

is_admin = False
if "ADMIN_EMAIL" in st.secrets and email == st.secrets["ADMIN_EMAIL"].lower().strip():
    is_admin = True
    st.session_state.role = "admin"
    nav_options.append("Admin Console")

nav_cols = st.columns(len(nav_options))
for i, option in enumerate(nav_options):
    bt_type = "primary" if st.session_state.current_page == option else "secondary"
    if nav_cols[i].button(option, type=bt_type, use_container_width=True):
        st.session_state.current_page = option
        st.rerun()

st.divider()
page = st.session_state.current_page

# --- PAGE: DASHBOARD ---
if page == "Dashboard":
    st.markdown(f"## Welcome, {st.session_state.name}")
    st.markdown("System Overview")
    
    res_saved = supabase.table('reading_list').select('id', count='exact').eq('user_email', email).execute()
    saved_count = res_saved.count if res_saved.count else 0
    
    res_notes = supabase.table('general_notes').select('id', count='exact').eq('user_email', email).execute()
    notes_count = res_notes.count if res_notes.count else 0
    
    res_pnotes = supabase.table('reading_list').select('id', count='exact').neq('notes', '').eq('user_email', email).execute()
    pnotes_count = res_pnotes.count if res_pnotes.count else 0

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Tracked Keywords", len(st.session_state.keywords))
    m2.metric("Saved Papers", saved_count)
    m3.metric("Literature Notes", pnotes_count)
    m4.metric("General Notes", notes_count)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    col_info, col_tools = st.columns([2, 1])
    with col_info:
        with st.container(border=True):
            st.markdown("### Workstation Protocol")
            st.markdown("""
            **1. Monitor:** Define tracking keywords in *Active Tracking* to receive continuous 48-hour literature updates from PubMed.  
            **2. Acquire:** Run deep searches in the *Literature Discovery* targeting specific filters or authors.  
            **3. Synthesize:** Save relevant papers to the *Reading Room* to write notes and further analyze them.  
            **4. Consolidate:** Review your collective notes and hypotheses in the *My Notebook*, and export a formatted APA manuscript directly to Word.
            """)
    with col_tools:
        with st.container(border=True):
            st.markdown("### External Research Tools")
            st.link_button("NotebookLM Access", "https://notebooklm.google.com/", use_container_width=True)
            st.link_button("Perplexity AI Engine", "https://www.perplexity.ai/", use_container_width=True)

# --- PAGE: ACTIVE TRACKING ---
elif page == "Active Tracking":
    st.markdown("## Active Tracking")
    st.markdown("Automated monitoring of relevant publications from the last 48 hours.")
    st.info("Need full text? Visit [Sci-Hub](https://sci-hub.se/) for manual access.")

    # Logic to add a new tracking filter to the user's profile
    def add_tracking_kw():
        new_kw = st.session_state.new_trk_kw.strip()
        if new_kw and new_kw not in st.session_state.keywords:
            st.session_state.keywords.append(new_kw)
            kw_str = ",".join(st.session_state.keywords)
            # Update Supabase database
            supabase.table('users').update({'keywords': kw_str}).eq('email', email).execute()
        st.session_state.new_trk_kw = ""

    # Input for new tracking filters
    with st.container(border=True):
        st.text_input("Add Tracking Filter (Press Enter):", key="new_trk_kw", on_change=add_tracking_kw, placeholder="Enter terminology (e.g., Default Mode Network)")

        # Display the active filters
        if st.session_state.keywords:
            st.caption("Active Filters:")
            
            # Start the horizontal flex container
            st.markdown('<div class="tag-container">', unsafe_allow_html=True)
            
            for kw in st.session_state.keywords:
                # Wrap each button in a tag-btn div
                st.markdown('<div class="tag-btn">', unsafe_allow_html=True)
                if st.button(f"{kw} ✖", key=f"del_trk_{kw}"):
                    st.session_state.keywords.remove(kw)
                    kw_str = ",".join(st.session_state.keywords)
                    supabase.table('users').update({'keywords': kw_str}).eq('email', email).execute()
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
                
            st.markdown('</div>', unsafe_allow_html=True) # Close tag-container

    # Fetch recently published articles
    recent = get_summaries(st.session_state.keywords, "", days=2)
    
    if not recent and st.session_state.keywords:
        st.info("No new publications identified within the last 48 hours.")
        
    # Render article cards
    for art in recent:
        with st.container(border=True):
            pid, ttl = str(art['Id']), art['Title']
            auths, jrnl, pdate = ", ".join(art['AuthorList']), art['Source'], art['PubDate']
            
            # Title row with Save/Remove button
            col_title, col_sv = st.columns([0.85, 0.15])
            with col_title:
                st.markdown(f"#### {ttl}")
            
            with col_sv:
                # Check database status for the save button
                check_sv = supabase.table('reading_list').select('id, notes').eq('pmid', pid).eq('user_email', email).execute()
                is_sv = bool(check_sv.data)
                has_notes = is_sv and bool(check_sv.data[0].get('notes', '').strip())
                
                if is_sv:
                    st.markdown('<div class="remove-btn-wrapper">', unsafe_allow_html=True)
                    if st.button("Remove", key=f"sv_trk_{pid}", use_container_width=True):
                        if has_notes:
                            st.session_state[f"confirm_rm_trk_{pid}"] = True
                        else:
                            toggle_reading_list(pid, ttl, jrnl, auths, pdate)
                            st.toast("Removed from Reading Room")
                            st.rerun()
                    st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="save-btn-wrapper">', unsafe_allow_html=True)
                    if st.button("Save", key=f"sv_trk_{pid}", use_container_width=True):
                        toggle_reading_list(pid, ttl, jrnl, auths, pdate)
                        st.toast("Article saved to Reading Room")
                        st.rerun()
                    st.markdown('</div>', unsafe_allow_html=True)

            # Note protection dialog
            if st.session_state.get(f"confirm_rm_trk_{pid}"):
                st.warning("This paper has notes. Removing it will delete them. Continue?")
                if st.button("Confirm Delete", key=f"conf_rm_trk_{pid}", type="primary"):
                    toggle_reading_list(pid, ttl, jrnl, auths, pdate)
                    del st.session_state[f"confirm_rm_trk_{pid}"]
                    st.rerun()

            st.markdown(f"<div class='paper-metadata'><b>Authors:</b> {auths}<br><b>Publication:</b> <i>{jrnl}</i> | <b>Date:</b> {pdate}</div>", unsafe_allow_html=True)
            
            # Action buttons
            c1, c2, c3 = st.columns(3)
            pub_url = f"https://pubmed.ncbi.nlm.nih.gov/{pid}/"
            ezproxy_url = f"https://ezproxy.haifa.ac.il/login?url={pub_url}"
            
            with c1: st.link_button("University Access (Haifa)", ezproxy_url, use_container_width=True)
            with c2: st.link_button("Web Access (General)", pub_url, use_container_width=True)
            with c3:
                if st.button("AI Summary", key=f"ai_sum_trk_{pid}", use_container_width=True): 
                    st.session_state[f"show_sum_trk_{pid}"] = True
            
            if st.session_state.get(f"show_sum_trk_{pid}"):
                with st.spinner("Analyzing abstract..."):
                    summary = generate_ai_summary(fetch_abstract(pid))
                    st.markdown(f"<div class='ai-summary-box'>{summary}</div>", unsafe_allow_html=True)
                if st.button("Dismiss Analysis", key=f"cls_trk_{pid}", type="secondary"): 
                    del st.session_state[f"show_sum_trk_{pid}"]
                    st.rerun()

# --- PAGE: Literature Discovery ---
elif page == "Literature Discovery":
    st.markdown("## Literature Discovery")
    st.info("Need full text? Visit [Sci-Hub](https://sci-hub.se/) for manual access.")

    # Logic to add search keywords to the current discovery session
    def add_discovery_kw():
        new_kw = st.session_state.new_disc_kw.strip()
        if new_kw and new_kw not in st.session_state.discovery_keywords:
            st.session_state.discovery_keywords.append(new_kw)
        st.session_state.new_disc_kw = ""

    # Search Configuration Container
    with st.container(border=True):
        st.markdown("### Filters")
        st.text_input("Keywords (Press Enter to Add):", key="new_disc_kw", on_change=add_discovery_kw)
        
        # Displaying active session keywords as horizontal tags
        if st.session_state.discovery_keywords:
            st.caption("Active Search Keywords:")
            # FIXED: Horizontal flexbox container to prevent vertical stacking
            st.markdown('<div class="tag-container">', unsafe_allow_html=True)
            for kw in st.session_state.discovery_keywords:
                # Wrap each button in a tag-btn div for CSS alignment
                st.markdown('<div class="tag-btn">', unsafe_allow_html=True)
                if st.button(f"{kw} ✖", key=f"del_disc_{kw}"):
                    st.session_state.discovery_keywords.remove(kw)
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        
        st.markdown("<br>", unsafe_allow_html=True)
        sauth = st.text_input("Target Author Name:", value="")
        
        # Search parameters: logic and timeframe
        c3, c4 = st.columns(2)
        with c3: slogic = st.radio("Search Logic (AND/OR):", ["OR", "AND"], horizontal=True)
        with c4: timeframe = st.selectbox("Time Range:", ["Week", "Month", "Year", "10 Years"], index=1)
        
        # Execution and Clearing controls
        col_exec, col_clear = st.columns([1, 1])
        with col_exec:
            if st.button("Execute Search", type="primary", use_container_width=True):
                st.session_state.discovery_ran = True
                t_map = {"Week": 7, "Month": 30, "Year": 365, "10 Years": 3650}
                with st.spinner("Analyzing PubMed results..."):
                    st.session_state.feed_results = get_summaries(st.session_state.discovery_keywords, sauth, t_map[timeframe], slogic)
        with col_clear:
            if st.button("Clear Search", type="secondary", use_container_width=True):
                st.session_state.discovery_keywords = []
                st.session_state.feed_results = []
                st.session_state.discovery_ran = False
                st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    # Rendering search results
    if st.session_state.discovery_ran:
        if not st.session_state.feed_results:
            st.warning("No articles found matching your criteria. Try adjusting your filters.")
        else:
            for art in st.session_state.feed_results:
                with st.container(border=True):
                    pid, ttl = str(art['Id']), art['Title']
                    auths, jrnl, pdate = ", ".join(art['AuthorList']), art['Source'], art['PubDate']
                    
                    # Row Layout: Title on the left, Save/Remove on the right
                    col_title, col_sv = st.columns([0.85, 0.15])
                    with col_title:
                        st.markdown(f"#### {ttl}")
                    
                    with col_sv:
                        # Check database for existing entries and saved notes
                        check_sv = supabase.table('reading_list').select('id, notes').eq('pmid', pid).eq('user_email', email).execute()
                        is_sv = bool(check_sv.data)
                        has_notes = is_sv and bool(check_sv.data[0].get('notes', '').strip())
                        
                        if is_sv:
                            # Apply remove-btn styling (Pastel Turquoise & No-Wrap)
                            st.markdown('<div class="remove-btn-wrapper">', unsafe_allow_html=True)
                            if st.button("Remove", key=f"sv_disc_{pid}", use_container_width=True):
                                if has_notes:
                                    # Flag for deletion confirmation if notes exist
                                    st.session_state[f"confirm_rm_disc_{pid}"] = True
                                else:
                                    toggle_reading_list(pid, ttl, jrnl, auths, pdate)
                                    st.toast("Removed from Reading Room")
                                    st.rerun()
                            st.markdown('</div>', unsafe_allow_html=True)
                        else:
                            # Apply save-btn styling (White background)
                            st.markdown('<div class="save-btn-wrapper">', unsafe_allow_html=True)
                            if st.button("Save", key=f"sv_disc_{pid}", use_container_width=True):
                                toggle_reading_list(pid, ttl, jrnl, auths, pdate)
                                st.toast("Article saved to Reading Room")
                                st.rerun()
                            st.markdown('</div>', unsafe_allow_html=True)

                    # Note Protection: Dialog for papers with existing content
                    if st.session_state.get(f"confirm_rm_disc_{pid}"):
                        st.warning("This paper has saved notes. Removing it will permanently delete them. Continue?")
                        if st.button("Confirm Delete", key=f"conf_rm_disc_{pid}", type="primary"):
                            toggle_reading_list(pid, ttl, jrnl, auths, pdate)
                            del st.session_state[f"confirm_rm_disc_{pid}"]
                            st.toast("Removed from Reading Room")
                            st.rerun()

                    # Article Metadata display
                    st.markdown(f"<div class='paper-metadata'><b>Authors:</b> {auths}<br><b>Publication:</b> <i>{jrnl}</i> | <b>Date:</b> {pdate}</div>", unsafe_allow_html=True)
                    
                    # Bottom Action Buttons
                    c1, c2, c3 = st.columns(3)
                    pub_url = f"https://pubmed.ncbi.nlm.nih.gov/{pid}/"
                    ezproxy_url = f"https://ezproxy.haifa.ac.il/login?url={pub_url}"
                    
                    with c1: st.link_button("University Access (Haifa)", ezproxy_url, use_container_width=True)
                    with c2: st.link_button("Web Access (General)", pub_url, use_container_width=True)
                    with c3:
                        if st.button("AI Summary", key=f"ai_disc_{pid}", use_container_width=True): 
                            st.session_state[f"show_sum_{pid}"] = True
                    
                    # AI Summary display container
                    if st.session_state.get(f"show_sum_{pid}"):
                        with st.spinner("Analyzing PubMed data..."):
                            summary = generate_ai_summary(fetch_abstract(pid))
                            st.markdown(f"<div class='ai-summary-box'>{summary}</div>", unsafe_allow_html=True)
                        if st.button("Dismiss Analysis", key=f"cls_s_{pid}", type="secondary"): 
                            del st.session_state[f"show_sum_{pid}"]
                            st.rerun()

# --- PAGE: READING ROOM ---
elif page == "Reading Room":
    st.markdown("## Reading Room")
    st.info("Need full text? Visit [Sci-Hub](https://sci-hub.se/) for manual access.")
    
    res = supabase.table('reading_list').select('*').eq('user_email', email).order('last_edited', desc=True).execute()
    items = res.data
    
    if not items:
        st.info("The reading room is currently empty. Transfer papers from Active Tracking or Literature Discovery to begin processing.")
        
    for item in items:
        pmid = item['pmid']
        with st.container(border=True):
            
            col_title, col_bm = st.columns([11, 1])
            with col_title:
                st.markdown(f"#### {item['title']}")
            with col_bm:
                st.markdown('<div class="remove-btn-marker"></div>', unsafe_allow_html=True)
                if st.button("Remove", key=f"rm_{pmid}", type="secondary"): 
                    toggle_reading_list(pmid, "", "", "", "")
                    st.toast("Article removed from Reading Room")
                    st.rerun()

            st.markdown(f"<div class='paper-metadata'><b>Authors:</b> {item['authors']}<br><b>Publication:</b> <i>{item['journal']}</i> | <b>Date:</b> {item['date']}</div>", unsafe_allow_html=True)
            
            c1, c2, c3 = st.columns(3)
            pub_url = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
            ezproxy_url = f"https://ezproxy.haifa.ac.il/login?url={pub_url}"
            q = urllib.parse.quote(f"Findings of: {item['title']}")
            
            with c1: st.link_button("University Access (Haifa)", ezproxy_url, use_container_width=True)
            with c2: st.link_button("Web Access (General)", pub_url, use_container_width=True)
            with c3: st.link_button("Investigate via Perplexity", f"https://www.perplexity.ai/search?q={q}", use_container_width=True)
            
            st.markdown("##### Analytical Notes")
            new_note = st.text_area("Write methodological insights, critiques, or hypotheses:", value=item.get('notes', ''), key=f"nt_rr_{pmid}", height=120, label_visibility="collapsed")
            if st.button("Save Notes", key=f"sv_rr_{pmid}", type="secondary"):
                supabase.table('reading_list').update({
                    'notes': new_note, 
                    'last_edited': datetime.now().strftime("%Y-%m-%d %H:%M")
                }).eq('pmid', pmid).eq('user_email', email).execute()
                
                st.toast("Notes saved.")
                st.rerun()

# --- PAGE: My NOTEBOOK ---
elif page == "My Notebook":
    st.markdown("## My Notebook")
    
    with st.container(border=True):
        st.download_button("Export to Docx", export_notes_to_word(), f"My_Notebook_{datetime.now().strftime('%Y%m%d')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", type="secondary")
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["General Research Notes & Thoughts", "Literature Notes"])
    
    with tab1:
        with st.container(border=True):
            st.markdown("### Add New Note")
            new_gen_c = st.text_area("Content:", key="new_gen_note", height=100, label_visibility="collapsed", placeholder="Enter your thoughts here...")
            if st.button("Append Note", type="primary"):
                if new_gen_c.strip():
                    data = {"user_email": email, "content": new_gen_c.strip(), "date": datetime.now().strftime("%Y-%m-%d %H:%M")}
                    supabase.table('general_notes').insert(data).execute()
                    st.rerun()
        
        st.markdown("<br>", unsafe_allow_html=True)
        res_g = supabase.table('general_notes').select('*').eq('user_email', email).order('id', desc=True).execute()
        
        for note in res_g.data:
            nid = note['id']
            with st.container(border=True):
                st.markdown(note['content'])
                c1, c2 = st.columns([5, 1])
                c1.caption(f"Added: {note['date']}")
                if c2.button("✖ Delete", key=f"del_g_{nid}", use_container_width=True, type="secondary"):
                    supabase.table('general_notes').delete().eq('id', nid).execute()
                    st.rerun()
                    
    with tab2:
        res_p = supabase.table('reading_list').select('*').neq('notes', '').eq('user_email', email).order('last_edited', desc=True).execute()
        p_notes = res_p.data
                
        if not p_notes:
            st.info("No active literature notes detected")
            
        for note in p_notes:
            pmid = note['pmid']
            with st.container(border=True):
                st.markdown(f"**{note['title']}**")
                st.markdown(note['notes'])
                
                c1, c2, c3 = st.columns([3, 2, 1])
                c1.caption(f"Last edited: {note['last_edited']}")
                
                pub_url = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/"
                
                c2.link_button("🔗 Source URL", pub_url, use_container_width=True)
                
                if c3.button("✖ Delete Note", key=f"del_p_{pmid}", use_container_width=True, type="secondary"):
                    supabase.table('reading_list').update({
                        'notes': '', 
                        'last_edited': datetime.now().strftime("%Y-%m-%d %H:%M")
                    }).eq('pmid', pmid).eq('user_email', email).execute()
                    st.rerun()

# --- PAGE: USER GUIDE ---
elif page == "User Guide":
    st.markdown("## Welcome to My Research Assistant!")
    st.markdown("""
    This workspace is dedicated to supporting the scientific journey by bridging the gap between publication and synthesis. It provides a focused environment for monitoring the latest literature, organizing discoveries, and documenting the evolution of research insights.
    By reducing logistical complexities and optimizing the research process, the platform facilitates the seamless integration of emerging scientific developments into daily academic practice.
    
    ### 1. Dashboard
    View your active metrics like tracked keywords, saved papers, and total independent ideas.

    ### 2. Active Literature Tracking
    **Purpose:** Stay updated automatically with the latest science that matches your interests.
    *   **Action:** Add specific keywords or terminology as "Tracking Filters". You can add multiple filters, and the system will use them to continuously monitor PubMed for new publications.
    *   **Result:** Every time you visit this page, the system securely connects to PubMed and retrieves all papers published in the **last 48 hours** matching your filters.
    *   **Workflow:** Read abstracts, Generate AI Summaries directly, and Save papers to your Reading Room.

    ### 3. Literature Discovery
    **Purpose:** Conduct deep historical searches for specific topics or authors.
    *   **Action:** Enter "Subject Keywords", specify an author if applicable, and adjust the historical range to query the PubMed database. Use AND/OR logic to refine your results.
    *   **Result:** Retrieves past papers matching your specific criteria. If you have tracking filters set in Active Literature Tracking, you can also apply those here to further narrow down results. If there are no results, try adjusting your filters or expanding the timeframe.
    *   **Workflow:** This is best for finding foundational papers or conducting targeted queries outside the 48-hour auto-tracking window. Search terms here are temporary and clear between sessions.

    ### 4. Reading Room
    **Purpose:** Save and organize papers for later review and note-taking.
    *   **Action:** Review all papers you've saved from Active Literature Tracking or Literature Discovery.
    *   **Result:** A list of your saved papers with metadata. You can write detailed analytical notes for each paper, which are stored in your personal profile. Use the "Investigate via Perplexity" button to explore complex topics mentioned in the paper through an external AI engine.
    *   **Workflow:** Write analytical notes for each paper directly in the interface. Saving notes attaches them securely to the paper and will also appear in the "Literature Notes" section in your notebook. Use the "Investigate via Perplexity" button for further investigation.

    ### 5. My Notebook
    **Purpose:** A centralized hub for all your research notes and thoughts.
    *   **General Notes:** A notepad for random hypotheses or general theoretical ideas.
    *   **Literature Notes:** A read-only view of all the notes you wrote in the Reading Room, displayed as easy-to-read cards.
    *   **Export to Docx:** Click this to generate a Word Document compiling all your reading list notes and general thoughts in a structured format.

    ### 6. Settings
    **Purpose:** Manage your profile and account settings.
    *   **Investigator Identity:** Update your name and view your registered email. (Email is fixed as it serves as your unique identifier in the system.)
    *   **Log Out:** Clear your session and log out of the system.
                
    """)

# --- PAGE: SETTINGS ---
elif page == "Settings":
    st.markdown("## User Settings")
    
    with st.container(border=True):
        st.markdown("### Investigator Identity")
        n_up = st.text_input("Researcher Name:", value=st.session_state.name)
        st.text_input("Email:", value=st.session_state.user_email, disabled=True)
        st.markdown("<br>", unsafe_allow_html=True)
        
        c1, c2 = st.columns([1, 8])
        with c1:
            if st.button("Save Changes", type="primary"):
                supabase.table('users').update({'name': n_up}).eq('email', email).execute()
                st.session_state.name = n_up
                st.success("Configuration saved.")
                st.rerun()
        
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("Log Out", type="secondary"):
        st.session_state.clear()
        st.rerun()

# --- PAGE: ADMIN CONSOLE (Exclusive View) ---
elif page == "Admin Console" and is_admin:
    st.markdown("## System Administration Console")
    st.markdown("Executive telemetry and user management.")
    
    c_users = supabase.table('users').select('email', count='exact').execute()
    c_docs = supabase.table('reading_list').select('id', count='exact').execute()
    c_logins = supabase.table('login_history').select('id', count='exact').execute()

    m1, m2, m3 = st.columns(3)
    m1.metric("Registered Investigators", c_users.count if c_users.count else 0)
    m2.metric("Total Papers Processed", c_docs.count if c_docs.count else 0)
    m3.metric("Total Authentication Events", c_logins.count if c_logins.count else 0)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    st.markdown("### User Management & Activity Tracking")
    all_users = supabase.table('users').select('*').order('created_at', desc=True).execute().data
    
    for u in all_users:
        u_email = u['email']
        with st.container(border=True):
            r_docs = supabase.table('reading_list').select('id', count='exact').eq('user_email', u_email).execute()
            doc_count = r_docs.count if r_docs.count else 0
            
            r_login = supabase.table('login_history').select('login_time').eq('user_email', u_email).order('login_time', desc=True).limit(1).execute()
            last_login = r_login.data[0]['login_time'][:16] if r_login.data else "Never"
            
            kws = u.get('keywords', '')
            kw_list =[k for k in kws.split(',') if k]
            
            st.markdown(f"**{u.get('name', 'Unnamed')}** (`{u_email}`)")
            col1, col2, col3, col4 = st.columns([1.5, 1.5, 3, 1.5])
            col1.write(f"**Saved Docs:** {doc_count}")
            col2.write(f"**Last Login:** {last_login}")
            
            with col3:
                if kw_list:
                    st.write("**Tracking:** " + ", ".join(kw_list))
                else:
                    st.write("**Tracking:** None")
            
            with col4:
                if st.button("Delete User", key=f"adm_del_{u_email}"):
                    st.session_state[f"confirm_delete_{u_email}"] = True
                
                if st.session_state.get(f"confirm_delete_{u_email}", False):
                    st.warning("Confirm Deletion?")
                    if st.button("Yes, Delete", key=f"confirm_yes_{u_email}", type="primary"):
                        supabase.table('reading_list').delete().eq('user_email', u_email).execute()
                        supabase.table('general_notes').delete().eq('user_email', u_email).execute()
                        supabase.table('login_history').delete().eq('user_email', u_email).execute()
                        supabase.table('users').delete().eq('email', u_email).execute()
                        st.success(f"User {u_email} deleted.")
                        st.session_state[f"confirm_delete_{u_email}"] = False
                        st.rerun()